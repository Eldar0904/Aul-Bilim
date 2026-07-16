#!/usr/bin/env python3
"""Export Russian site copy to a readable Word document."""
from __future__ import annotations

import json
import re
from datetime import date
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REGISTRY = ROOT / "uploads" / "copy-registry.js"
OUT = ROOT / "Aul-Bilim-Russian-Copy.docx"
PAGES = [
    ("index.html", "Главная"),
    ("about.html", "О нас"),
    ("programs.html", "Программы"),
    ("school.html", "Страница школы"),
]


class RuSpanParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.entries: list[tuple[str, str]] = []
        self._in_ru = False
        self._key = ""
        self._buf: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag != "span":
            return
        ad = {k: v for k, v in attrs if v is not None}
        if ad.get("lang") == "ru":
            self._in_ru = True
            self._key = ad.get("data-copy", "")
            self._buf = []

    def handle_endtag(self, tag: str) -> None:
        if tag == "span" and self._in_ru:
            text = re.sub(r"\s+", " ", "".join(self._buf)).strip()
            if text:
                self.entries.append((self._key, text))
            self._in_ru = False
            self._key = ""
            self._buf = []

    def handle_data(self, data: str) -> None:
        if self._in_ru:
            self._buf.append(data)


def load_registry() -> list[dict]:
    src = REGISTRY.read_text(encoding="utf-8")
    m = re.search(r"window\.COPY_REGISTRY\s*=\s*(\[.*?\]);", src, re.S)
    if not m:
        raise SystemExit("COPY_REGISTRY not found")
    return json.loads(m.group(1))


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Aul Bilim")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(27, 79, 138)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Русский текст сайта")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(232, 130, 10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Экспорт: {date.today().isoformat()}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Документ содержит весь русский UI-текст публичных страниц. "
        "Источник: index.html, about.html, programs.html, school.html."
    )
    note.runs[0].font.size = Pt(10)
    note.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_copy_item(doc: Document, text: str, key: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    if key:
        p.add_run(f"  [{key}]").font.size = Pt(8)


def unique_entries(entries: list[tuple[str, str]]) -> list[tuple[str, str]]:
    seen: set[str] = set()
    out: list[tuple[str, str]] = []
    for key, text in entries:
        if text in seen:
            continue
        seen.add(text)
        out.append((key, text))
    return out


def add_page_titles(doc: Document) -> None:
    add_heading(doc, "Заголовки страниц (title)", 1)
    for fname, label in PAGES:
        html = (ROOT / fname).read_text(encoding="utf-8")
        tm = re.search(r'<title[^>]*data-ru="([^"]+)"', html)
        if tm:
            p = doc.add_paragraph()
            p.add_run(f"{label}").bold = True
            p.add_run(f" — {tm.group(1)}")
    doc.add_paragraph()


def add_html_copy(doc: Document) -> None:
    add_heading(doc, "Текст по страницам", 1)
    for fname, label in PAGES:
        parser = RuSpanParser()
        parser.feed((ROOT / fname).read_text(encoding="utf-8"))
        entries = unique_entries(parser.entries)
        if not entries:
            continue
        add_heading(doc, label, 2)
        for key, text in entries:
            add_copy_item(doc, text, key)
        doc.add_paragraph()


def add_dynamic_copy(doc: Document) -> None:
    add_heading(doc, "Карта и динамический интерфейс", 1)
    map_src = (ROOT / "map.js").read_text(encoding="utf-8")
    strings: list[str] = []
    for m in re.finditer(r"bi\('([^']*)',\s*'([^']*)'\)", map_src):
        strings.append(m.group(2))
    for m in re.finditer(r"ru:\s*'([^']+)'", map_src):
        val = m.group(1)
        if len(val) > 12:
            strings.append(val)
    seen: set[str] = set()
    for text in strings:
        if text in seen:
            continue
        seen.add(text)
        add_copy_item(doc, text)


def add_registry_appendix(doc: Document) -> None:
    reg = load_registry()
    ru_reg = [e for e in reg if e.get("key", "").endswith("-ru")]
    add_heading(doc, "Приложение: CMS-ключи", 1)
    doc.add_paragraph(
        "Ключи для админки (copy-registry.js). Основной текст выше — полные строки из HTML.",
        style="Intense Quote",
    )
    by_page: dict[str, list[dict]] = {}
    for e in ru_reg:
        by_page.setdefault(e.get("page", "unknown"), []).append(e)
    order = ["site_shared", "index", "about", "programs", "school"]
    page_labels = {
        "site_shared": "Общее (нав / footer)",
        "index": "index.html",
        "about": "about.html",
        "programs": "programs.html",
        "school": "school.html",
    }
    for page in order + sorted(k for k in by_page if k not in order):
        items = by_page.get(page)
        if not items:
            continue
        add_heading(doc, page_labels.get(page, page), 2)
        for e in items:
            label = " ".join((e.get("label") or "").split())
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(e["key"]).font.size = Pt(9)
            p.add_run(f" — {label}").font.size = Pt(10)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_page(doc)
    add_page_titles(doc)
    add_html_copy(doc)
    add_dynamic_copy(doc)
    add_registry_appendix(doc)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
